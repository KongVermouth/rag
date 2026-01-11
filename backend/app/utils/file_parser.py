"""
文件解析工具 - 将不同格式文件转换为文本
基于 src/file_utils.py 的 FileParser 类
"""
import logging
import traceback
import re
from pathlib import Path
from typing import BinaryIO, Optional

import fitz as pymupdf  # PyMuPDF < 1.24.0 使用 fitz 导入
import docx
import html2text
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class FileParser:
    """单文件解析器：将不同格式文件转换为纯文本"""

    def __init__(self):
        # 配置 HTML 转 Markdown 的参数
        self.h2t = html2text.HTML2Text()
        self.h2t.ignore_links = False
        self.h2t.ignore_images = True
        self.h2t.body_width = 0
        self.h2t.ignore_emphasis = False

    def parse_file(self, file_path: Path) -> str:
        """根据文件后缀解析文件"""
        suffix = file_path.suffix.lower()
        try:
            if suffix == '.docx':
                return self._parse_docx(file_path)
            elif suffix == '.pdf':
                return self._parse_pdf(file_path)
            elif suffix == '.html':
                return self._parse_html(file_path)
            elif suffix == '.md':
                return self._parse_md(file_path)
            elif suffix == '.txt':
                return self._parse_txt(file_path)
            else:
                logger.warning(f"不支持的文件格式: {file_path}")
                return ""
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"解析文件失败 {file_path}:\n{error_detail}")
            raise

    def _parse_docx(self, file_path: Path) -> str:
        """解析 .docx 文件，保留标题层级和表格"""
        doc = docx.Document(file_path)
        md_lines = []

        p_idx = 0
        t_idx = 0

        for elem in doc.element.body:
            if elem.tag.endswith('p'):
                if p_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[p_idx]
                    p_idx += 1

                    text = paragraph.text.strip()
                    if not text:
                        continue

                    style_name = paragraph.style.name
                    if 'Heading' in style_name:
                        try:
                            level = int(style_name.split(' ')[-1])
                        except ValueError:
                            level = 1
                        md_lines.append(f"{'#' * level} {text}\n")
                    else:
                        md_lines.append(f"{text}\n")

            elif elem.tag.endswith('tbl'):
                if t_idx < len(doc.tables):
                    table = doc.tables[t_idx]
                    t_idx += 1
                    md_table = self._docx_table_to_md(table)
                    md_lines.append(md_table)
                    md_lines.append("\n")

        return "".join(md_lines)

    def _docx_table_to_md(self, table) -> str:
        """将 Word 表格转换为 Markdown 表格"""
        md_table = []
        if not table.rows:
            return ""

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        md_table.append("| " + " | ".join(headers) + " |")
        md_table.append("| " + " | ".join(["---"] * len(headers)) + " |")

        for row in table.rows[1:]:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            md_table.append("| " + " | ".join(row_data) + " |")

        return "\n".join(md_table) + "\n"

    def _parse_pdf(self, file_path: Path) -> str:
        """解析文字版 PDF"""
        doc = pymupdf.open(file_path)
        md_content = []

        for page in doc:
            html_content = page.get_text("html")
            md_text = self.h2t.handle(html_content)
            md_content.append(md_text)

        doc.close()

        full_text = "\n\n".join(md_content)
        return full_text

    def _parse_html(self, file_path: Path) -> str:
        """解析 .html 文件"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            html_content = f.read()

        if not html_content.strip():
            return ""

        soup = BeautifulSoup(html_content, 'html.parser')

        for script in soup(["script", "style", "noscript", "header", "footer"]):
            script.decompose()

        block_tags = {'p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                      'li', 'ul', 'ol', 'tr', 'td', 'th', 'br', 'hr',
                      'article', 'section', 'aside', 'main', 'nav'}

        text_parts = []

        def extract_text_recursive(element):
            for child in element.children:
                if isinstance(child, str):
                    text = child.strip()
                    if text:
                        text_parts.append(text)
                else:
                    tag_name = child.name
                    extract_text_recursive(child)
                    if tag_name in block_tags:
                        text_parts.append("\n")

        root = soup.body if soup.body else soup
        extract_text_recursive(root)

        raw_text = "".join(text_parts)

        if not raw_text.strip():
            raw_text = soup.get_text(separator='\n')

        clean_text = re.sub(r'\n{3,}', '\n\n', raw_text)
        return clean_text

    def _parse_md(self, file_path: Path) -> str:
        """读取 .md 文件"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        return content

    def _parse_txt(self, file_path: Path) -> str:
        """读取 .txt 文件"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()


# 全局文件解析器实例
file_parser = FileParser()
